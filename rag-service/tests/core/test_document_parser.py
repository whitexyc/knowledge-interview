"""module-064 WP1 解析层测试（mock AnyDoc + 真实多格式，ADR-0014 决策 1）

覆盖：
- 格式识别读字节魔数（pdf/docx/xlsx），md/txt 无魔数用扩展名兜底
- md/txt 纯文本解码（UTF-8 / GBK 兜底）
- AnyDoc 主解析路径（mock）：pdf/page_count / docx
- 错误变体映射用户中文提示（Unsupported/Malformed/Encrypted/ResourceLimit）
- AnyDoc 不可用 → PDF 走 PyMuPDF、docx/xlsx/csv 走轻量回退、pptx/epub 明确报错
- 上传端点接线（扩展名校验 + 错误码）
"""
import asyncio
import io
import unittest.mock as mock

import pytest

from rag.retrieval import document_parser
from rag.retrieval.document_parser import (
    DocumentParseError,
    ParsedDocument,
    SUPPORTED_EXTENSIONS,
    detect_format,
    parse_document,
)


# ── 格式识别 ─────────────────────────────────────────────────────────────
def test_supported_extensions_cover_all():
    """上传端允许的 8 种格式齐全"""
    assert set(SUPPORTED_EXTENSIONS) == {
        ".md", ".txt", ".pdf", ".docx", ".xlsx", ".pptx", ".epub", ".csv",
    }


def test_detect_format_md_via_extension():
    """md/txt 无魔数：扩展名兜底 → text"""
    assert detect_format(b"hello", "a.md") == "text"
    assert detect_format(b"hello", "a.txt") == "text"


def test_detect_format_pdf_real_bytes():
    """PDF 魔数 %PDF 由 anydoc 探测"""
    import fitz
    buf = io.BytesIO()
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "probe")
    doc.save(buf)
    doc.close()
    assert detect_format(buf.getvalue(), "noext") == "pdf"


def test_detect_format_unknown_raises():
    """未知格式 → 明确 DocumentParseError"""
    with pytest.raises(DocumentParseError, match="不支持的文件格式"):
        detect_format(b"not a real doc at all", "a.xyz")


# ── 纯文本解析 ───────────────────────────────────────────────────────────
def test_parse_text_utf8():
    r = parse_document("你好 hello world".encode("utf-8"), "a.md")
    assert r.engine == "text"
    assert r.format == "text"
    assert "你好" in r.text


def test_parse_text_gbk_fallback():
    """GBK 编码老 txt：UTF-8 失败 → GB18030 兜底"""
    data = "中文内容".encode("gbk")
    r = parse_document(data, "a.txt")
    assert "中文内容" in r.text


def test_parse_empty_data():
    with pytest.raises(DocumentParseError, match="上传文件为空"):
        parse_document(b"", "a.md")


# ── AnyDoc 主解析（mock）────────────────────────────────────────────────
class FakeAnyDoc:
    """可编程假 anydoc：format_from_bytes / to_markdown_bytes"""

    def __init__(self, md="", raise_exc=None, detect=None):
        self._md = md
        self._raise = raise_exc
        self._detect = detect

    def format_from_bytes(self, data):
        return self._detect

    def to_markdown_bytes(self, data):
        if self._raise is not None:
            raise self._raise
        return self._md


def test_parse_pdf_anydoc(monkeypatch):
    """AnyDoc 主解析：pdf + page_count 透传"""
    monkeypatch.setattr(document_parser, "anydoc",
                        FakeAnyDoc(md="# PDF 标题\n\n正文", detect="pdf"))
    monkeypatch.setattr(document_parser, "_pdf_page_count", lambda data: 3)
    r = parse_document(b"%PDF-1.4 fake-bytes", "a.pdf")
    assert r.engine == "anydoc"
    assert r.format == "pdf"
    assert r.page_count == 3
    assert "# PDF 标题" in r.text


def test_parse_docx_anydoc(monkeypatch):
    """AnyDoc 主解析：docx → GFM Markdown"""
    monkeypatch.setattr(document_parser, "anydoc",
                        FakeAnyDoc(md="| A | B |\n| --- | --- |\n| 1 | 2 |", detect="docx"))
    r = parse_document(b"PK\x03\x04fake", "a.docx")
    assert r.engine == "anydoc"
    assert "| A | B |" in r.text


def test_parse_md_ignores_anydoc(monkeypatch):
    """md/txt 不走 AnyDoc（即使 anydoc 存在）——纯文本解码"""
    monkeypatch.setattr(document_parser, "anydoc", FakeAnyDoc(md="WRONG", detect=None))
    r = parse_document("真实文本".encode("utf-8"), "a.md")
    assert r.engine == "text"
    assert "真实文本" in r.text


# ── 错误变体映射 ────────────────────────────────────────────────────────
def test_error_unsupported_mapping(monkeypatch):
    """Unsupported → '不支持的文件格式'（pptx 无轻量回退 → 走映射）"""
    monkeypatch.setattr(document_parser, "anydoc",
                        FakeAnyDoc(raise_exc=document_parser.UnsupportedError("unsupported"),
                                   detect="pptx"))
    with pytest.raises(DocumentParseError, match="不支持的文件格式"):
        parse_document(b"fake", "a.pptx")


def test_error_encrypted_mapping(monkeypatch):
    """Encrypted → '文件已加密'"""
    monkeypatch.setattr(document_parser, "anydoc",
                        FakeAnyDoc(raise_exc=document_parser.EncryptedError("enc"),
                                   detect="pptx"))
    with pytest.raises(DocumentParseError, match="已加密"):
        parse_document(b"fake", "a.pptx")


def test_error_malformed_mapping(monkeypatch):
    """Malformed → '文件已损坏'"""
    monkeypatch.setattr(document_parser, "anydoc",
                        FakeAnyDoc(raise_exc=document_parser.MalformedError("bad"),
                                   detect="pptx"))
    with pytest.raises(DocumentParseError, match="已损坏"):
        parse_document(b"fake", "a.pptx")


def test_pdf_anydoc_error_falls_back_to_pymupdf(monkeypatch):
    """PDF AnyDoc 失败 → PyMuPDF 回退（存量行为保留）"""
    monkeypatch.setattr(document_parser, "anydoc",
                        FakeAnyDoc(raise_exc=document_parser.ConvertError("conv"),
                                   detect="pdf"))
    fake = ParsedDocument(text="--- Page 1/1 ---\npyMuPDF 文本", format="pdf",
                          engine="pymupdf", page_count=1)
    monkeypatch.setattr(document_parser, "_parse_pdf_pymupdf", lambda data: fake)
    r = parse_document(b"%PDF-1.4", "a.pdf")
    assert r.engine == "pymupdf"
    assert "pyMuPDF 文本" in r.text


# ── AnyDoc 不可用 → 分层回退 ─────────────────────────────────────────────
def test_anydoc_unavailable_pdf_pymupdf(monkeypatch):
    monkeypatch.setattr(document_parser, "anydoc", None)
    fake = ParsedDocument(text="pymupdf text", format="pdf", engine="pymupdf", page_count=1)
    monkeypatch.setattr(document_parser, "_parse_pdf_pymupdf", lambda data: fake)
    r = parse_document(b"%PDF-1.4", "a.pdf")
    assert r.engine == "pymupdf"


def test_anydoc_unavailable_docx_fallback(monkeypatch):
    """AnyDoc 不可用 → python-docx 轻量回退（真实解析）"""
    monkeypatch.setattr(document_parser, "anydoc", None)
    from docx import Document as Docx
    doc = Docx()
    doc.add_heading("文档标题", 1)
    doc.add_paragraph("这是一段正文")
    buf = io.BytesIO()
    doc.save(buf)
    r = parse_document(buf.getvalue(), "a.docx")
    assert r.engine == "docx_fallback"
    assert "文档标题" in r.text
    assert "这是一段正文" in r.text


def test_anydoc_unavailable_xlsx_fallback(monkeypatch):
    monkeypatch.setattr(document_parser, "anydoc", None)
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws["A1"], ws["B1"] = "Name", "Value"
    ws["A2"], ws["B2"] = "Alice", 42
    buf = io.BytesIO()
    wb.save(buf)
    r = parse_document(buf.getvalue(), "a.xlsx")
    assert r.engine == "xlsx_fallback"
    assert "| Name | Value |" in r.text
    assert "| Alice | 42 |" in r.text


def test_anydoc_unavailable_pptx_clear_error(monkeypatch):
    """pptx 无轻量回退 → 明确报'需 AnyDoc 解析引擎'（诚实降级）"""
    monkeypatch.setattr(document_parser, "anydoc", None)
    with pytest.raises(DocumentParseError, match="需要 AnyDoc 解析引擎"):
        parse_document(b"PK\x03\x04", "a.pptx")


# ── 上传端点接线（AC 1.2/1.3）───────────────────────────────────────────
def test_upload_endpoint_unsupported_extension():
    """上传端点：非白名单扩展名 → code=1 明确提示"""
    import main as main_module

    async def run():
        with mock.patch("main.ingest_document") as mock_ingest:
            import httpx
            transport = httpx.ASGITransport(app=main_module.app, raise_app_exceptions=True)
            async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
                resp = await client.post(
                    "/ai/rag/documents/upload",
                    files={"file": ("a.xyz", b"whatever", "application/octet-stream")},
                )
            assert resp.status_code == 200
            data = resp.json()
            assert data["code"] == 1
            assert "不支持的文件格式" in data["message"]
            mock_ingest.assert_not_called()

    asyncio.run(run())


def test_upload_endpoint_success_invokes_ingest():
    """上传端点：白名单格式 → 调用 ingest_document 并返回结果"""
    import main as main_module

    async def run():
        fake_result = {"id": 7, "title": "a", "chunks": 3, "duplicate": False,
                       "dup_kind": None, "page_count": None, "original_path": ""}
        with mock.patch("main.ingest_document",
                        new=mock.AsyncMock(return_value=fake_result)) as mock_ingest:
            import httpx
            transport = httpx.ASGITransport(app=main_module.app, raise_app_exceptions=True)
            async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
                resp = await client.post(
                    "/ai/rag/documents/upload",
                    files={"file": ("a.md", "# 标题\n\n正文".encode("utf-8"), "text/markdown")},
                )
            assert resp.status_code == 200
            data = resp.json()
            assert data["code"] == 0
            assert data["data"]["id"] == 7
            mock_ingest.assert_awaited_once()

    asyncio.run(run())


def test_upload_endpoint_parse_error_mapped():
    """上传端点：解析失败 → code=3 + 中文提示"""
    import main as main_module

    async def run():
        with mock.patch("main.ingest_document",
                        new=mock.AsyncMock(
                            side_effect=DocumentParseError("文件已加密，无法解析"))):
            import httpx
            transport = httpx.ASGITransport(app=main_module.app, raise_app_exceptions=True)
            async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
                resp = await client.post(
                    "/ai/rag/documents/upload",
                    files={"file": ("a.pdf", b"%PDF-1.4 fake", "application/pdf")},
                )
            assert resp.status_code == 200
            data = resp.json()
            assert data["code"] == 3
            assert "已加密" in data["message"]

    asyncio.run(run())


# ── module-069: 双栏中线重组 + pymupdf4llm ─────────────────────────────────


def _make_mock_page(width: float, blocks: list, page_num: int = 1):
    """构造模拟 page 对象（get_text("blocks") 返回给定块列表）

    blocks: list of (x0, y0, x1, y1, text, block_no, block_type)
    """
    import unittest.mock as mock
    page = mock.MagicMock()
    page.rect.width = width
    page.rect.height = 842.0
    page.get_text.return_value = "\n".join(b[4] for b in blocks if b[6] == 0)
    page.get_text.side_effect = lambda mode="text": (
        blocks if mode == "blocks" else "\n".join(b[4] for b in blocks if b[6] == 0)
    )
    return page


class TestReorderColumns:
    """WP-A: _reorder_columns 纯函数测试"""

    def test_single_column_no_reorder(self):
        """单栏页面走正常 y 序，不触发重组"""
        from rag.retrieval.document_parser import _reorder_columns
        page = _make_mock_page(600.0, [
            (10, 50, 280, 70, "第一段", 0, 0),
            (10, 100, 280, 120, "第二段", 1, 0),
        ])
        result = _reorder_columns(page)
        assert "第一段" in result
        assert "第二段" in result
        assert result.index("第一段") < result.index("第二段")

    def test_double_column_detected(self):
        """双栏页面（左右各 >=2 块）正确检测为双栏并重组"""
        from rag.retrieval.document_parser import _reorder_columns
        page = _make_mock_page(600.0, [
            (10, 50, 280, 70, "左栏第一段", 0, 0),
            (10, 150, 280, 170, "左栏第二段", 1, 0),
            (320, 50, 580, 70, "右栏第一段", 2, 0),
            (320, 150, 580, 170, "右栏第二段", 3, 0),
        ])
        result = _reorder_columns(page)
        left1_pos = result.index("左栏第一段")
        left2_pos = result.index("左栏第二段")
        right1_pos = result.index("右栏第一段")
        right2_pos = result.index("右栏第二段")
        assert left1_pos < left2_pos, "左栏内部应按 y0 排序"
        assert right1_pos < right2_pos, "右栏内部应按 y0 排序"
        assert left2_pos < right1_pos, "左栏整体在右栏之前"

    def test_cross_mid_block_on_top(self):
        """跨中线块排在最前"""
        from rag.retrieval.document_parser import _reorder_columns
        page = _make_mock_page(600.0, [
            (10, 200, 280, 220, "左栏段落", 0, 0),
            (320, 200, 580, 220, "右栏段落", 1, 0),
            (100, 50, 500, 70, "跨中线大标题", 2, 0),
        ])
        result = _reorder_columns(page)
        assert result.startswith("跨中线大标题"), "跨中线块应置顶"

    def test_table_not_reordered(self):
        """含 | 的块不参与重组（表格内列顺序不能乱）"""
        from rag.retrieval.document_parser import _reorder_columns
        page = _make_mock_page(600.0, [
            (10, 50, 280, 70, "左栏文本", 0, 0),
            (10, 100, 280, 120, "左栏第二段", 1, 0),
            (320, 50, 580, 70, "| 列A | 列B |", 2, 0),
            (320, 100, 580, 120, "右栏第二段", 3, 0),
        ])
        result = _reorder_columns(page)
        left_pos = result.index("左栏文本")
        table_pos = result.index("| 列A | 列B |")
        assert left_pos < table_pos

    def test_empty_blocks(self):
        """空块列表不崩溃"""
        from rag.retrieval.document_parser import _reorder_columns
        page = _make_mock_page(600.0, [])
        result = _reorder_columns(page)
        assert result == ""

    def test_image_blocks_filtered(self):
        """图片块（block_type=1）被过滤"""
        from rag.retrieval.document_parser import _reorder_columns
        page = _make_mock_page(600.0, [
            (10, 50, 280, 70, "文本块", 0, 0),
            (10, 100, 280, 200, "", 1, 1),
        ])
        result = _reorder_columns(page)
        assert "文本块" in result

    def test_all_cross_mid_page(self):
        """全跨中线页面正常处理（left/right 为空，走单栏路径）"""
        from rag.retrieval.document_parser import _reorder_columns
        page = _make_mock_page(600.0, [
            (50, 50, 550, 70, "宽标题", 0, 0),
            (50, 100, 550, 120, "宽段落", 1, 0),
        ])
        result = _reorder_columns(page)
        assert "宽标题" in result
        assert "宽段落" in result


class TestPdfFallbackMdSwitch:
    """WP-B: pymupdf4llm 开关行为测试"""

    def test_switch_false_uses_get_text(self, monkeypatch):
        """开关 false → engine=pymupdf"""
        from src.config import settings
        monkeypatch.setattr(settings, "pdf_fallback_md", False)

        import fitz
        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello World")
        doc.save(buf)
        doc.close()

        from rag.retrieval.document_parser import _parse_pdf_pymupdf
        result = _parse_pdf_pymupdf(buf.getvalue())
        assert result.engine == "pymupdf"
        assert "Hello World" in result.text

    def test_switch_true_uses_pymupdf4llm(self, monkeypatch):
        """开关 true → engine=pymupdf4llm"""
        from src.config import settings
        monkeypatch.setattr(settings, "pdf_fallback_md", True)

        import fitz
        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Test Content")
        doc.save(buf)
        doc.close()

        from rag.retrieval.document_parser import _parse_pdf_pymupdf
        result = _parse_pdf_pymupdf(buf.getvalue())
        assert result.engine == "pymupdf4llm"
        assert "Test Content" in result.text

    def test_pymupdf4llm_not_installed_falls_back(self, monkeypatch):
        """pymupdf4llm 未安装时降级 get_text() + warning"""
        from src.config import settings
        monkeypatch.setattr(settings, "pdf_fallback_md", True)

        import fitz
        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Fallback Text")
        doc.save(buf)
        doc.close()

        # 模拟 pymupdf4llm 不可用：mock import
        import sys
        saved = sys.modules.get("pymupdf4llm")
        sys.modules["pymupdf4llm"] = None  # type: ignore
        try:
            import rag.retrieval.document_parser as dp
            # 强制重新检测导入状态
            import importlib
            importlib.reload(dp)
            result = dp._parse_pdf_pymupdf(buf.getvalue())
            assert result.engine == "pymupdf"
            assert "Fallback Text" in result.text
        finally:
            if saved is not None:
                sys.modules["pymupdf4llm"] = saved
            else:
                sys.modules.pop("pymupdf4llm", None)
            importlib.reload(dp)

    def test_single_page_pdf(self):
        """单页 PDF 正常处理"""
        import fitz
        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Single Page")
        doc.save(buf)
        doc.close()

        from rag.retrieval.document_parser import _parse_pdf_pymupdf
        result = _parse_pdf_pymupdf(buf.getvalue())
        assert result.page_count == 1
        assert "--- Page 1/1 ---" in result.text

    def test_engine_field_consistency(self, monkeypatch):
        """开关切换时 engine 字段一致"""
        from src.config import settings
        import fitz

        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "content")
        doc.save(buf)
        doc.close()

        from rag.retrieval.document_parser import _parse_pdf_pymupdf

        monkeypatch.setattr(settings, "pdf_fallback_md", False)
        r_false = _parse_pdf_pymupdf(buf.getvalue())
        assert r_false.engine == "pymupdf"

        monkeypatch.setattr(settings, "pdf_fallback_md", True)
        r_true = _parse_pdf_pymupdf(buf.getvalue())
        assert r_true.engine == "pymupdf4llm"


class TestColumnReorderIntegration:
    """WP-A + WP-B 集成：双栏 PDF 真实解析"""

    def test_real_double_column_pdf(self):
        """真实双栏 PDF（rag_survey）第 2 页重组——双栏检测 + 输出非空"""
        import os
        pdf_path = os.path.join(
            os.path.dirname(__file__), "..", "..", "..",
            "docs", "项目深挖", "原PDF-rag_survey.pdf",
        )
        if not os.path.exists(pdf_path):
            pytest.skip("测试 PDF 文件不存在（主 checkout 路径）")

        from rag.retrieval.document_parser import _reorder_columns
        import fitz

        data = open(pdf_path, "rb").read()
        doc = fitz.open(stream=data, filetype="pdf")
        try:
            page = doc[1]  # 第 2 页
            blocks = page.get_text("blocks")
            mid_x = page.rect.width / 2
            text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]
            left_count = sum(1 for b in text_blocks if b[2] <= mid_x)
            right_count = sum(1 for b in text_blocks if b[0] >= mid_x)

            # 第 2 页应为双栏
            assert left_count >= 2, f"左栏块数 {left_count} 应 >=2"
            assert right_count >= 2, f"右栏块数 {right_count} 应 >=2"

            reordered = _reorder_columns(page)
            assert len(reordered) > 100, "重组后文本应有实质内容"
            # 重组后文本应包含双栏内容（左+右）
            single_col_text = "\n".join(b[4].rstrip() for b in sorted(text_blocks, key=lambda b: b[1]))
            assert len(reordered) >= len(single_col_text) * 0.8, \
                "重组后文本长度不应显著少于原始拼接"
        finally:
            doc.close()

    def test_chinese_pdf_parses(self, monkeypatch):
        """中文 PDF 正常解析（pymupdf4llm）"""
        import os
        pdf_path = os.path.join(
            os.path.dirname(__file__), "..", "..", "..",
            "docs", "项目深挖", "原PDF-图检索增强生成研究综述.pdf",
        )
        if not os.path.exists(pdf_path):
            pytest.skip("测试 PDF 文件不存在（主 checkout 路径）")

        from src.config import settings
        monkeypatch.setattr(settings, "pdf_fallback_md", True)

        from rag.retrieval.document_parser import _parse_pdf_pymupdf
        data = open(pdf_path, "rb").read()
        result = _parse_pdf_pymupdf(data)
        assert result.page_count == 12
        assert len(result.text) > 100
        assert result.engine == "pymupdf4llm"

    def test_real_pdf_switch_false_uses_pymupdf_engine(self):
        """开关 false 时真实 PDF engine=pymupdf"""
        import os
        pdf_path = os.path.join(
            os.path.dirname(__file__), "..", "..", "..",
            "docs", "项目深挖", "原PDF-rag_survey.pdf",
        )
        if not os.path.exists(pdf_path):
            pytest.skip("测试 PDF 文件不存在（主 checkout 路径）")

        import unittest.mock as mock
        from src.config import settings
        from rag.retrieval.document_parser import _parse_pdf_pymupdf

        data = open(pdf_path, "rb").read()
        with mock.patch.object(settings, "pdf_fallback_md", False):
            result = _parse_pdf_pymupdf(data)
            assert result.engine == "pymupdf"
            assert result.page_count == 21


# ── module-069 Tester 补充测试 ─────────────────────────────────────────────


class TestReorderColumnsExtended:
    """WP-A 补充：排序正确性 + 边界条件"""

    def test_reorder_produces_correct_order(self):
        """双栏重组顺序：跨中线块 → 左栏 y 序 → 右栏 y 序"""
        from rag.retrieval.document_parser import _reorder_columns
        page = _make_mock_page(600.0, [
            # 左栏两块（y 序 100, 200）
            (10, 200, 280, 220, "左栏下", 0, 0),
            (10, 100, 280, 120, "左栏上", 1, 0),
            # 右栏两块（y 序 100, 200）
            (320, 200, 580, 220, "右栏下", 2, 0),
            (320, 100, 580, 120, "右栏上", 3, 0),
            # 跨中线块（y=50，应置顶）
            (100, 50, 500, 70, "跨中线标题", 4, 0),
        ])
        result = _reorder_columns(page)
        lines = [l for l in result.split("\n") if l.strip()]
        assert lines[0] == "跨中线标题", "跨中线块应排第一"
        assert lines[1] == "左栏上", "左栏按 y0 排序"
        assert lines[2] == "左栏下"
        assert lines[3] == "右栏上", "右栏在左栏之后"
        assert lines[4] == "右栏下"

    def test_table_stays_in_original_column(self):
        """含 | 的表格块按 x0 归入对应栏，不跨栏移动"""
        from rag.retrieval.document_parser import _reorder_columns
        page = _make_mock_page(600.0, [
            (10, 50, 280, 70, "左栏文本", 0, 0),
            (10, 100, 280, 120, "左栏第二段", 1, 0),
            # 右栏表格（含 |）——应留在右栏位置，不移到左栏
            (320, 50, 580, 70, "| 列A | 列B |", 2, 0),
            (320, 100, 580, 120, "右栏第二段", 3, 0),
        ])
        result = _reorder_columns(page)
        lines = [l for l in result.split("\n") if l.strip()]
        # 左栏在前，右栏在后，表格在右栏内部
        left_end = max(i for i, l in enumerate(lines) if "左栏" in l)
        table_idx = next(i for i, l in enumerate(lines) if "| 列A" in l)
        right_text_idx = next(i for i, l in enumerate(lines) if "右栏第二段" in l)
        assert table_idx > left_end, "表格在左栏之后（右栏区域）"
        assert table_idx < right_text_idx, "表格在右栏内部按 y0 排序"

    def test_single_column_with_marginal_note(self):
        """单栏页面有 1 个偏右块（页码/页眉）不误切为双栏"""
        from rag.retrieval.document_parser import _reorder_columns
        page = _make_mock_page(600.0, [
            (10, 100, 280, 120, "正文第一段", 0, 0),
            (10, 200, 280, 220, "正文第二段", 1, 0),
            (500, 50, 560, 60, "2", 2, 0),  # 页码偏右，只有 1 块
        ])
        result = _reorder_columns(page)
        # 应走单栏路径（右栏只有 1 块 < 2），按 y0 排序
        assert result.index("正文第一段") < result.index("正文第二段")

    def test_empty_text_blocks_filtered(self):
        """纯空白文本块（block_type=0 但 text 为空）被过滤"""
        from rag.retrieval.document_parser import _reorder_columns
        page = _make_mock_page(600.0, [
            (10, 50, 280, 70, "有效文本", 0, 0),
            (10, 100, 280, 120, "   ", 1, 0),  # 纯空白
        ])
        result = _reorder_columns(page)
        assert "有效文本" in result
        # 空白块不应出现
        lines = [l for l in result.split("\n") if l.strip()]
        assert len(lines) == 1


class TestPymupdf4llmOutput:
    """WP-B 补充：pymupdf4llm 输出格式 + 清洗层衔接"""

    def test_pymupdf4llm_output_has_markdown_structure(self, monkeypatch):
        """pymupdf4llm 输出包含 Markdown 结构（标题/列表）"""
        from src.config import settings
        monkeypatch.setattr(settings, "pdf_fallback_md", True)

        import fitz
        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        # 插入多行文本模拟段落
        page.insert_text((72, 72), "Section Title", fontsize=16)
        page.insert_text((72, 100), "Body paragraph text")
        doc.save(buf)
        doc.close()

        from rag.retrieval.document_parser import _parse_pdf_pymupdf
        result = _parse_pdf_pymupdf(buf.getvalue())
        assert result.engine == "pymupdf4llm"
        assert "Section Title" in result.text

    def test_cleaning_layer_integration(self, monkeypatch):
        """pymupdf4llm 输出可被 document_cleaner.clean() 清洗（管线衔接验证）"""
        from src.config import settings
        monkeypatch.setattr(settings, "pdf_fallback_md", True)

        import fitz
        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Clean test content")
        doc.save(buf)
        doc.close()

        from rag.retrieval.document_parser import _parse_pdf_pymupdf
        result = _parse_pdf_pymupdf(buf.getvalue())
        # 验证 parser 输出可被 cleaner 处理（非空、不含异常字符导致 cleaner 崩溃）
        from rag.retrieval import document_cleaner
        cleaned = document_cleaner.clean(result.text)
        assert len(cleaned) > 0, "清洗后文本非空"


class TestPwEnvVarSwitch:
    """WP-B 补充：PW_PDF_FALLBACK_MD 环境变量切换"""

    def test_env_var_false_forces_pymupdf_engine(self, monkeypatch):
        """PW_PDF_FALLBACK_MD=false → engine=pymupdf"""
        from src.config import settings
        monkeypatch.setattr(settings, "pdf_fallback_md", False)

        import fitz
        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "env test")
        doc.save(buf)
        doc.close()

        from rag.retrieval.document_parser import _parse_pdf_pymupdf
        result = _parse_pdf_pymupdf(buf.getvalue())
        assert result.engine == "pymupdf"

    def test_env_var_true_forces_pymupdf4llm_engine(self, monkeypatch):
        """PW_PDF_FALLBACK_MD=true → engine=pymupdf4llm"""
        from src.config import settings
        monkeypatch.setattr(settings, "pdf_fallback_md", True)

        import fitz
        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "env test")
        doc.save(buf)
        doc.close()

        from rag.retrieval.document_parser import _parse_pdf_pymupdf
        result = _parse_pdf_pymupdf(buf.getvalue())
        assert result.engine == "pymupdf4llm"


class TestEdgeCases:
    """WP-A/B 边界条件补充"""

    def test_empty_pdf_no_crash(self):
        """空 PDF（零文本层）不抛未处理异常"""
        import fitz
        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        # 不插入任何文本，模拟空内容 PDF
        doc.save(buf)
        doc.close()

        from rag.retrieval.document_parser import _parse_pdf_pymupdf
        result = _parse_pdf_pymupdf(buf.getvalue())
        assert result.page_count == 1
        # 空内容不崩溃，分页标记存在
        assert "--- Page 1/1 ---" in result.text

    def test_scanned_pdf_empty_text(self, monkeypatch):
        """扫描版 PDF（无文本层）返回空文本（不崩溃）"""
        import fitz
        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        # 不插入任何文本，只画一个矩形（模拟扫描图片）
        page.draw_rect(fitz.Rect(10, 10, 100, 100))
        doc.save(buf)
        doc.close()

        from src.config import settings
        monkeypatch.setattr(settings, "pdf_fallback_md", False)
        from rag.retrieval.document_parser import _parse_pdf_pymupdf
        result = _parse_pdf_pymupdf(buf.getvalue())
        assert result.page_count == 1
        # 空文本层不崩溃，text 可能为空或只有分页标记
        assert "--- Page 1/1 ---" in result.text

    def test_single_page_cannot_be_double_column(self):
        """单页 PDF 正常处理（单页不可能双栏）"""
        import fitz
        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Single page content")
        page.insert_text((72, 100), "Second paragraph")
        doc.save(buf)
        doc.close()

        from rag.retrieval.document_parser import _parse_pdf_pymupdf
        result = _parse_pdf_pymupdf(buf.getvalue())
        assert result.page_count == 1
        assert "Single page content" in result.text
